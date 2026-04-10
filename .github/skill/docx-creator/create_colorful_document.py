#!/usr/bin/env python3
"""
Microsoft DOCX Colorful Creator - Enterprise Professional Document Generator
Software GBB Americas | 4-Color Microsoft Logo Palette

Author: paulasilva@microsoft.com
Team: Software GBB Americas

Creates vibrant, enterprise-grade Word documents with the 4-color Microsoft
logo palette (Blue, Green, Yellow, Red), rotating colors across sections.

Features:
- Modern themed cover page with 4-color bar
- Automatic Table of Contents (always included)
- Document properties auto-filled (author, title, subject, keywords)
- Header: "Software GBB Americas" branding
- Rich visual elements: tables, icons, highlights, callouts, diagrams
- Standard closing page with contact info and next steps
- All content in English
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from datetime import datetime
import re


# =============================================================================
# MICROSOFT LOGO COLORS (4-color palette)
# =============================================================================
class MSLogoColors:
    """Microsoft 4-color logo palette + supporting neutrals."""
    # Logo colors
    BLUE = RGBColor(0x00, 0xA4, 0xEF)       # #00A4EF
    GREEN = RGBColor(0x7F, 0xBA, 0x00)       # #7FBA00
    YELLOW = RGBColor(0xFF, 0xB9, 0x00)      # #FFB900
    RED = RGBColor(0xF2, 0x50, 0x22)         # #F25022

    # Logo hex strings (for cell backgrounds)
    BLUE_HEX = "00A4EF"
    GREEN_HEX = "7FBA00"
    YELLOW_HEX = "FFB900"
    RED_HEX = "F25022"

    # Supporting colors
    DARK_BLUE = RGBColor(0x00, 0x78, 0xD4)   # #0078D4 (MS primary blue)
    DARK_GREEN = RGBColor(0x10, 0x7C, 0x10)  # #107C10
    DARK_RED = RGBColor(0xD1, 0x34, 0x38)    # #D13438

    # Neutrals
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    BLACK = RGBColor(0x00, 0x00, 0x00)
    DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
    MEDIUM_GRAY = RGBColor(0x50, 0x50, 0x50)
    GRAY = RGBColor(0x66, 0x66, 0x66)
    LIGHT_GRAY = RGBColor(0xF3, 0xF3, 0xF3)
    ALT_ROW = RGBColor(0xF9, 0xF9, 0xF9)

    # Light tints for highlight boxes
    LIGHT_BLUE = "DAEEF7"
    LIGHT_GREEN = "E8F5D0"
    LIGHT_YELLOW = "FFF4D6"
    LIGHT_RED = "FDE0DA"

    # Color rotation order
    ROTATION = ["blue", "green", "yellow", "red"]

    @classmethod
    def get(cls, name: str) -> RGBColor:
        """Get RGBColor by name."""
        mapping = {
            "blue": cls.BLUE,
            "green": cls.GREEN,
            "yellow": cls.YELLOW,
            "red": cls.RED,
        }
        return mapping.get(name.lower(), cls.BLUE)

    @classmethod
    def get_hex(cls, name: str) -> str:
        """Get hex string by name."""
        mapping = {
            "blue": cls.BLUE_HEX,
            "green": cls.GREEN_HEX,
            "yellow": cls.YELLOW_HEX,
            "red": cls.RED_HEX,
        }
        return mapping.get(name.lower(), cls.BLUE_HEX)

    @classmethod
    def get_light(cls, name: str) -> str:
        """Get light tint hex by name (for highlight boxes)."""
        mapping = {
            "blue": cls.LIGHT_BLUE,
            "green": cls.LIGHT_GREEN,
            "yellow": cls.LIGHT_YELLOW,
            "red": cls.LIGHT_RED,
        }
        return mapping.get(name.lower(), cls.LIGHT_BLUE)

    @classmethod
    def text_color(cls, name: str) -> RGBColor:
        """Get text color for headers (black for yellow, white for others)."""
        if name.lower() == "yellow":
            return cls.BLACK
        return cls.WHITE


# =============================================================================
# ICONS / UNICODE SYMBOLS
# =============================================================================
class Icons:
    """Unicode icons for document elements."""
    CHECK = "✅"
    WARNING = "⚠️"
    INFO = "ℹ️"
    IMPORTANT = "❗"
    TIP = "💡"
    PIN = "📌"
    ROCKET = "🚀"
    TARGET = "🎯"
    CHART = "📊"
    GEAR = "⚙️"
    LINK = "🔗"
    CALENDAR = "📅"
    PERSON = "👤"
    TEAM = "👥"
    EMAIL = "📧"
    PHONE = "📞"
    STAR = "⭐"
    CLOCK = "⏱️"
    SHIELD = "🛡️"
    KEY = "🔑"
    ARROW_RIGHT = "➡️"
    PROGRESS = "🔄"
    SUCCESS = "✅"
    NEXT = "▶️"


# =============================================================================
# COLORFUL DOCUMENT CREATOR CLASS
# =============================================================================
class MSColorfulDocCreator:
    """
    Enterprise-grade Microsoft Word document creator using the 4-color
    Microsoft logo palette. Produces professional, visually rich documents
    with modern cover page, TOC, color-coded sections, rich visual elements,
    and standard closing page.

    Always includes:
    - Cover page (themed, modern, with 4-color bar)
    - Table of Contents
    - Document properties (author, title, subject, keywords)
    - Header: Software GBB Americas
    - Footer: Microsoft Confidential | Page X of Y
    - Standard closing page (contact + next steps)
    - All content in English
    """

    # Default logo path (relative to skill assets)
    DEFAULT_LOGO = None  # Set at runtime via _resolve_logo_path()

    def __init__(
        self,
        title: str,
        subtitle: str = "",
        version: str = "1.0",
        client: str = "",
        author: str = "paulasilva@microsoft.com",
        date: str = None,
        team: str = "Software GBB Americas",
        info_lines: list = None,
        subject: str = "",
        keywords: str = "",
        category: str = "Technical Documentation",
        logo_path: str = None,
    ):
        """
        Initialize a new colorful enterprise document.

        Args:
            title: Document main title
            subtitle: Document subtitle/description
            version: Version number (e.g., "1.0")
            client: Client name (optional)
            author: Author email (default: paulasilva@microsoft.com)
            date: Document date (defaults to today)
            team: Team name (default: Software GBB Americas)
            info_lines: Optional list of extra info lines for cover page
            subject: Document subject for properties
            keywords: Document keywords for properties
            category: Document category for properties
            logo_path: Path to logo image for cover page (auto-resolves default)
        """
        self.doc = Document()
        self.title = title
        self.subtitle = subtitle
        self.version = version
        self.client = client
        self.author = author
        self.date = date or datetime.now().strftime("%B %d, %Y")
        self.team = team
        self.info_lines = info_lines or []
        self.subject = subject or subtitle
        self.keywords = keywords
        self.category = category
        self.logo_path = logo_path or self._resolve_logo_path()

        # Color rotation tracker
        self._color_index = 0

        # Setup
        self._setup_document_properties()
        self._setup_styles()
        self._setup_page_layout()

    @staticmethod
    def _resolve_logo_path():
        """Auto-resolve the default logo from the skill's assets folder."""
        import os
        # Try common skill paths (prefer white-bg version for documents)
        candidates = [
            os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                         'assets', 'logo-msft-github-white-bg.png'),
            os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                         'assets', 'logo-msft-github-color-black.png'),
            '/mnt/skills/user/ms-docx-colorful/assets/logo-msft-github-white-bg.png',
            '/mnt/skills/user/ms-docx-colorful/assets/logo-msft-github-color-black.png',
            '/mnt/user-data/uploads/logo-msft-github-color-black.png',
        ]
        for path in candidates:
            if os.path.exists(path):
                return path
        return None

    # =========================================================================
    # DOCUMENT SETUP
    # =========================================================================
    def _setup_document_properties(self):
        """Set document metadata/properties based on theme and subject."""
        core = self.doc.core_properties
        core.author = self.author
        core.title = self.title
        core.subject = self.subject
        core.keywords = self.keywords or f"{self.team}, Microsoft, {self.title}, Enterprise"
        core.category = self.category
        core.comments = f"Created by {self.team} | {self.author}"
        core.language = "en-US"

    def _setup_page_layout(self):
        """Configure page size and margins."""
        section = self.doc.sections[0]
        section.page_width = Inches(8.27)   # A4
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    def _setup_styles(self):
        """Configure all document styles with Segoe UI font family."""
        styles = self.doc.styles

        # Title style
        self._ensure_style(styles, 'MS Title', WD_STYLE_TYPE.PARAGRAPH,
                           font_name='Segoe UI Light', size=32, bold=False,
                           color=MSLogoColors.DARK_GRAY,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=0, space_after=6, line_spacing=1.0)

        # Subtitle style
        self._ensure_style(styles, 'MS Subtitle', WD_STYLE_TYPE.PARAGRAPH,
                           font_name='Segoe UI Semilight', size=16, bold=False,
                           color=MSLogoColors.GRAY,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=0, space_after=12, line_spacing=1.15)

        # Heading 1
        h1 = styles['Heading 1']
        self._apply_font(h1, 'Segoe UI', 18, True, MSLogoColors.BLUE)
        h1.paragraph_format.space_before = Pt(20)
        h1.paragraph_format.space_after = Pt(8)
        h1.paragraph_format.line_spacing = 1.15
        h1.paragraph_format.keep_with_next = True

        # Heading 2
        h2 = styles['Heading 2']
        self._apply_font(h2, 'Segoe UI', 15, True, MSLogoColors.DARK_BLUE)
        h2.paragraph_format.space_before = Pt(16)
        h2.paragraph_format.space_after = Pt(6)
        h2.paragraph_format.line_spacing = 1.15
        h2.paragraph_format.keep_with_next = True

        # Heading 3
        h3 = styles['Heading 3']
        self._apply_font(h3, 'Segoe UI Semibold', 13, True, MSLogoColors.DARK_GRAY)
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)
        h3.paragraph_format.line_spacing = 1.15
        h3.paragraph_format.keep_with_next = True

        # Heading 4
        h4 = self._ensure_heading(styles, 'Heading 4')
        self._apply_font(h4, 'Segoe UI', 12, True, MSLogoColors.MEDIUM_GRAY)
        h4.paragraph_format.space_before = Pt(10)
        h4.paragraph_format.space_after = Pt(4)
        h4.paragraph_format.line_spacing = 1.15

        # Normal body text
        normal = styles['Normal']
        self._apply_font(normal, 'Segoe UI', 11, False, MSLogoColors.BLACK)
        normal.paragraph_format.space_before = Pt(6)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.15

        # Caption
        self._ensure_style(styles, 'MS Caption', WD_STYLE_TYPE.PARAGRAPH,
                           font_name='Segoe UI', size=9, italic=True,
                           color=MSLogoColors.GRAY,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           space_before=4, space_after=8)

        # Source citation
        self._ensure_style(styles, 'MS Source', WD_STYLE_TYPE.PARAGRAPH,
                           font_name='Segoe UI', size=9, italic=True,
                           color=MSLogoColors.DARK_BLUE,
                           space_before=2, space_after=8)

        # List bullet
        lb = styles['List Bullet']
        self._apply_font(lb, 'Segoe UI', 11, False, MSLogoColors.BLACK)
        lb.paragraph_format.space_before = Pt(3)
        lb.paragraph_format.space_after = Pt(3)
        lb.paragraph_format.line_spacing = 1.15
        lb.paragraph_format.left_indent = Inches(0.5)

    def _ensure_style(self, styles, name, style_type, font_name='Segoe UI',
                      size=11, bold=False, italic=False, color=None,
                      alignment=None, space_before=6, space_after=8,
                      line_spacing=1.15):
        """Create or get a style and apply formatting."""
        if name not in [s.name for s in styles]:
            style = styles.add_style(name, style_type)
        else:
            style = styles[name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        if color:
            style.font.color.rgb = color
        if alignment is not None:
            style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.line_spacing = line_spacing
        return style

    def _ensure_heading(self, styles, name):
        """Get or create heading style."""
        if name in [s.name for s in styles]:
            return styles[name]
        return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    def _apply_font(self, style, name, size, bold, color):
        """Apply font settings to a style."""
        style.font.name = name
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color

    # =========================================================================
    # COLOR ROTATION
    # =========================================================================
    def next_color(self) -> str:
        """Get next color in rotation (blue → green → yellow → red → blue...)."""
        color = MSLogoColors.ROTATION[self._color_index % 4]
        self._color_index += 1
        return color

    def reset_color_rotation(self):
        """Reset color rotation to start."""
        self._color_index = 0

    # =========================================================================
    # COVER PAGE (Modern, themed, with logo)
    # =========================================================================
    def create_cover_page(self):
        """
        Create a modern, themed cover page with:
        - 4-color bar at top
        - Microsoft + GitHub logo (auto-resolved from assets)
        - Large title in dark gray
        - Subtitle
        - Document metadata
        - Confidentiality notice
        - 4-color bar at bottom
        """
        import os

        # 4-color bar at top of cover
        self._add_cover_color_bar()

        # Spacer
        for _ in range(2):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        # Logo (Microsoft + GitHub)
        if self.logo_path and os.path.exists(self.logo_path):
            logo_para = self.doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.paragraph_format.space_before = Pt(0)
            logo_para.paragraph_format.space_after = Pt(12)
            run = logo_para.add_run()
            run.add_picture(self.logo_path, width=Inches(3.5))

        # Spacer
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # Title
        title_para = self.doc.add_paragraph(self.title, style='MS Title')

        # Blue accent line under title
        self._add_accent_line(title_para, MSLogoColors.BLUE_HEX, thickness=12)

        # Subtitle
        if self.subtitle:
            sub_para = self.doc.add_paragraph(self.subtitle, style='MS Subtitle')

        # Spacer
        for _ in range(4):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        # Metadata block
        meta_items = []
        if self.client:
            meta_items.append(f"{Icons.TARGET}  Client: {self.client}")
        meta_items.append(f"{Icons.CALENDAR}  Date: {self.date}")
        meta_items.append(f"{Icons.GEAR}  Version: {self.version}")
        meta_items.append(f"{Icons.TEAM}  Team: {self.team}")
        meta_items.append(f"{Icons.EMAIL}  Author: {self.author}")

        # Add custom info lines
        for line in self.info_lines:
            meta_items.append(f"     {line}")

        for item in meta_items:
            para = self.doc.add_paragraph()
            run = para.add_run(item)
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
            run.font.color.rgb = MSLogoColors.DARK_GRAY
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.left_indent = Inches(0.5)

        # Spacer
        for _ in range(3):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        # Confidentiality notice
        conf_para = self.doc.add_paragraph()
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = conf_para.add_run("Microsoft Confidential")
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = MSLogoColors.DARK_RED

        # Bottom 4-color bar
        self._add_cover_color_bar()

        self.doc.add_page_break()

    def _add_cover_color_bar(self):
        """Add a 4-color horizontal bar using a single-row table."""
        table = self.doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        colors = [MSLogoColors.RED_HEX, MSLogoColors.GREEN_HEX,
                  MSLogoColors.BLUE_HEX, MSLogoColors.YELLOW_HEX]

        for i, color in enumerate(colors):
            cell = table.rows[0].cells[i]
            cell.text = ""
            self._set_cell_background(cell, color)
            # Set cell height (small bar)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            # Set preferred width
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '2000')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

        # Remove table borders for clean look
        self._remove_table_borders(table)

        # Set row height
        row = table.rows[0]
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '200')
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

    def _add_accent_line(self, paragraph, color_hex, thickness=8):
        """Add a colored bottom border to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(thickness))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # =========================================================================
    # TABLE OF CONTENTS (Always included)
    # =========================================================================
    def add_table_of_contents(self):
        """Add a Table of Contents. Always call after cover page."""
        toc_title = self.doc.add_paragraph("Table of Contents")
        toc_title.style = 'Heading 1'

        # TOC field
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()

        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        # Instruction
        note = self.doc.add_paragraph(
            "[Right-click and select 'Update Field' to refresh Table of Contents]"
        )
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in note.runs:
            r.font.name = 'Segoe UI'
            r.font.size = Pt(9)
            r.font.italic = True
            r.font.color.rgb = MSLogoColors.GRAY

        self.doc.add_page_break()

    # =========================================================================
    # HEADER / FOOTER
    # =========================================================================
    def setup_header_footer(self, header_text: str = None):
        """
        Configure header and footer.
        Header L1: {Title} | Software GBB Americas
        Header L2: Version {x.xx} | paulasilva@microsoft.com
        Footer: Microsoft Confidential | Page X of Y
        """
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True

        # HEADER
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        line1_text = header_text or f"{self.title} | {self.team}"
        run1 = header_para.add_run(line1_text)
        run1.font.name = 'Segoe UI'
        run1.font.size = Pt(9)
        run1.font.color.rgb = MSLogoColors.DARK_GRAY

        header_para.add_run().add_break()

        line2_text = f"Version {self.version} | {self.author}"
        run2 = header_para.add_run(line2_text)
        run2.font.name = 'Segoe UI'
        run2.font.size = Pt(8)
        run2.font.color.rgb = MSLogoColors.GRAY

        # Header bottom border in blue
        self._add_accent_line(header_para, MSLogoColors.BLUE_HEX, thickness=6)

        # FOOTER
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        conf_run = footer_para.add_run("Microsoft Confidential | ")
        conf_run.font.name = 'Segoe UI'
        conf_run.font.size = Pt(8)
        conf_run.font.color.rgb = MSLogoColors.GRAY

        self._add_page_number(footer_para)

    # =========================================================================
    # SECTION HEADINGS (Colored)
    # =========================================================================
    def add_section_heading(self, text: str, level: int = 1, color: str = None):
        """
        Add a colored section heading.

        Args:
            text: Heading text
            level: 1, 2, 3, or 4
            color: "blue", "green", "yellow", "red" or None for auto-rotation
        """
        if color is None:
            color = self.next_color()

        style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3', 4: 'Heading 4'}
        style = style_map.get(level, 'Heading 1')

        para = self.doc.add_paragraph(text, style=style)

        # Override color
        for run in para.runs:
            run.font.color.rgb = MSLogoColors.get(color)
            if color == "yellow":
                # Yellow headings get a darker shade for readability
                run.font.color.rgb = RGBColor(0xE0, 0xA0, 0x00)

        return para

    # =========================================================================
    # BODY TEXT
    # =========================================================================
    def add_body(self, text: str):
        """Add a body paragraph."""
        para = self.doc.add_paragraph(text, style='Normal')
        return para

    def add_bold_body(self, label: str, text: str):
        """Add a paragraph with bold label followed by normal text."""
        para = self.doc.add_paragraph()
        run_bold = para.add_run(label)
        run_bold.font.name = 'Segoe UI'
        run_bold.font.size = Pt(11)
        run_bold.font.bold = True
        run_bold.font.color.rgb = MSLogoColors.DARK_GRAY

        run_text = para.add_run(text)
        run_text.font.name = 'Segoe UI'
        run_text.font.size = Pt(11)
        return para

    # =========================================================================
    # LISTS
    # =========================================================================
    def add_bullets(self, items: list, level: int = 0):
        """Add bulleted list items."""
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25 + (level * 0.25))

    def add_numbered(self, items: list):
        """Add numbered list items."""
        for i, item in enumerate(items, 1):
            para = self.doc.add_paragraph(f"{i}. {item}", style='Normal')
            para.paragraph_format.left_indent = Inches(0.25)

    def add_icon_list(self, items: list):
        """
        Add a list with icons.
        Each item is a tuple: (icon, text) e.g. (Icons.CHECK, "Task completed")
        """
        for icon, text in items:
            para = self.doc.add_paragraph()
            run = para.add_run(f"{icon}  {text}")
            run.font.name = 'Segoe UI'
            run.font.size = Pt(11)
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)

    # =========================================================================
    # TABLES (Colored headers)
    # =========================================================================
    def add_table(self, headers: list, rows: list, color: str = None,
                  col_widths: list = None):
        """
        Add a formatted table with colored header row.

        Args:
            headers: List of column header strings
            rows: List of row data (each row is a list of strings)
            color: Header color ("blue", "green", "yellow", "red") or auto
            col_widths: Optional list of column widths in inches
        """
        if color is None:
            color = MSLogoColors.ROTATION[self._color_index % 4]

        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths if provided
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)

        # Header row
        header_cells = table.rows[0].cells
        text_color = MSLogoColors.text_color(color)
        for i, header in enumerate(headers):
            header_cells[i].text = header
            para = header_cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = 'Segoe UI'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = text_color
            self._set_cell_background(header_cells[i], MSLogoColors.get_hex(color))

        # Data rows with alternating shading
        for row_idx, row_data in enumerate(rows):
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                row.cells[i].text = str(cell_text)
                para = row.cells[i].paragraphs[0]
                for run in para.runs:
                    run.font.name = 'Segoe UI'
                    run.font.size = Pt(10)
                # Alternating row shading
                if row_idx % 2 == 1:
                    self._set_cell_background(row.cells[i], 'F3F3F3')

        self.doc.add_paragraph()  # Spacing
        return table

    # =========================================================================
    # VISUAL ELEMENTS
    # =========================================================================
    def add_color_bar(self):
        """Add a 4-color decorative separator bar between sections."""
        table = self.doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        colors = [MSLogoColors.RED_HEX, MSLogoColors.GREEN_HEX,
                  MSLogoColors.BLUE_HEX, MSLogoColors.YELLOW_HEX]

        for i, c in enumerate(colors):
            cell = table.rows[0].cells[i]
            cell.text = ""
            self._set_cell_background(cell, c)

        self._remove_table_borders(table)

        row = table.rows[0]
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '120')
        trHeight.set(qn('w:hRule'), 'exact')
        trPr.append(trHeight)

        self.doc.add_paragraph()  # spacing

    def add_highlight_box(self, title: str, text: str, color: str = "blue",
                          icon: str = None):
        """
        Add a colored highlight/callout box.

        Args:
            title: Bold title text
            text: Body text
            color: Box color
            icon: Optional icon prefix
        """
        # Create a single-cell table for the box
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]

        # Background tint
        self._set_cell_background(cell, MSLogoColors.get_light(color))

        # Left border color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '24')
        left.set(qn('w:color'), MSLogoColors.get_hex(color))
        left.set(qn('w:space'), '0')
        tcBorders.append(left)
        tcPr.append(tcBorders)

        # Title paragraph
        title_para = cell.paragraphs[0]
        prefix = f"{icon}  " if icon else ""
        run_title = title_para.add_run(f"{prefix}{title}")
        run_title.font.name = 'Segoe UI'
        run_title.font.size = Pt(11)
        run_title.font.bold = True
        run_title.font.color.rgb = MSLogoColors.get(color)

        # Body paragraph
        body_para = cell.add_paragraph()
        run_body = body_para.add_run(text)
        run_body.font.name = 'Segoe UI'
        run_body.font.size = Pt(10)
        run_body.font.color.rgb = MSLogoColors.DARK_GRAY

        self.doc.add_paragraph()  # spacing
        return table

    def add_note(self, title: str, text: str, color: str = "blue"):
        """Add a colored inline note (icon + bold prefix + text)."""
        icon_map = {
            "blue": Icons.INFO,
            "green": Icons.TIP,
            "yellow": Icons.WARNING,
            "red": Icons.IMPORTANT,
        }
        para = self.doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)

        prefix_run = para.add_run(f"{icon_map.get(color, Icons.INFO)} {title}")
        prefix_run.font.name = 'Segoe UI'
        prefix_run.font.size = Pt(11)
        prefix_run.font.bold = True
        prefix_run.font.color.rgb = MSLogoColors.get(color)

        text_run = para.add_run(f" {text}")
        text_run.font.name = 'Segoe UI'
        text_run.font.size = Pt(11)
        return para

    def add_key_takeaway(self, text: str, color: str = "blue"):
        """Add a Key Takeaway highlight box."""
        return self.add_highlight_box(
            "Key Takeaway", text, color=color, icon=Icons.KEY
        )

    def add_best_practice(self, text: str, color: str = "green"):
        """Add a Best Practice highlight box."""
        return self.add_highlight_box(
            "Best Practice", text, color=color, icon=Icons.STAR
        )

    def add_diagram_placeholder(self, description: str, color: str = "blue"):
        """Add a placeholder for a diagram/flow/chart."""
        return self.add_highlight_box(
            f"{Icons.CHART} Diagram", description, color=color
        )

    def add_code_block(self, code: str, language: str = ""):
        """Add a code block with monospace font."""
        if language:
            lang_para = self.doc.add_paragraph()
            run = lang_para.add_run(f"  {language}")
            run.font.name = 'Segoe UI'
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = MSLogoColors.GRAY

        para = self.doc.add_paragraph(code)
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)
        for run in para.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        return para

    def add_source(self, source_text: str, url: str = ""):
        """Add a source citation."""
        text = f"{Icons.LINK} Source: {source_text}"
        if url:
            text += f" ({url})"
        para = self.doc.add_paragraph(text, style='MS Source')
        return para

    # =========================================================================
    # VERSION HISTORY TABLE
    # =========================================================================
    def add_version_history(self, versions: list = None):
        """
        Add a version history table.

        Args:
            versions: List of dicts with keys: version, date, author, changes
                     If None, creates initial entry.
        """
        self.add_section_heading("Version History", level=2, color="blue")

        if versions is None:
            versions = [{
                "version": self.version,
                "date": self.date,
                "author": self.author,
                "changes": "Initial version"
            }]

        headers = ["Version", "Date", "Author", "Changes"]
        rows = [[v["version"], v["date"], v["author"], v["changes"]]
                for v in versions]
        self.add_table(headers, rows, color="blue")

    # =========================================================================
    # EXECUTIVE SUMMARY SECTION
    # =========================================================================
    def add_executive_summary(self, summary_text: str, color: str = "blue"):
        """Add an Executive Summary section with highlight box."""
        self.add_section_heading("Executive Summary", level=1, color=color)
        self.add_highlight_box("Overview", summary_text, color=color, icon=Icons.TARGET)

    # =========================================================================
    # GLOSSARY / ACRONYMS
    # =========================================================================
    def add_glossary(self, terms: list, color: str = "green"):
        """
        Add a Glossary/Acronyms section.

        Args:
            terms: List of tuples (term, definition)
            color: Table header color
        """
        self.add_section_heading("Glossary & Acronyms", level=1, color=color)
        headers = ["Term", "Definition"]
        rows = [[t, d] for t, d in terms]
        self.add_table(headers, rows, color=color)

    # =========================================================================
    # CLOSING PAGE (Standard)
    # =========================================================================
    def add_closing_page(self, next_steps: list = None,
                         contact_name: str = "Paula Silva",
                         contact_email: str = "paulasilva@microsoft.com",
                         contact_role: str = "Senior Software Engineer | Global Black Belt",
                         contact_team: str = "Software GBB Americas",
                         github_usernames: list = None,
                         additional_contacts: list = None):
        """
        Add the standard closing page with contact info and next steps.

        Args:
            next_steps: List of next step strings (optional)
            contact_name: Primary contact name
            contact_email: Primary contact email
            contact_role: Primary contact role
            contact_team: Contact team
            github_usernames: List of GitHub usernames (e.g. ["@paulanunes85", "@paulasilvatech"])
            additional_contacts: List of dicts with name, email, role, github (optional)
        """
        if github_usernames is None:
            github_usernames = ["@paulanunes85", "@paulasilvatech"]
        self.doc.add_page_break()

        # 4-color bar
        self.add_color_bar()

        # Thank You / Questions section
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(0)

        thank_para = self.doc.add_paragraph()
        thank_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = thank_para.add_run("Thank You")
        run.font.name = 'Segoe UI Light'
        run.font.size = Pt(28)
        run.font.color.rgb = MSLogoColors.DARK_BLUE

        questions_para = self.doc.add_paragraph()
        questions_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = questions_para.add_run("Questions? Let's connect.")
        run.font.name = 'Segoe UI'
        run.font.size = Pt(14)
        run.font.color.rgb = MSLogoColors.GRAY

        # Spacer
        self.doc.add_paragraph()

        # Next Steps (if provided)
        if next_steps:
            ns_heading = self.doc.add_paragraph()
            ns_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = ns_heading.add_run(f"{Icons.ROCKET}  Next Steps")
            run.font.name = 'Segoe UI'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = MSLogoColors.GREEN

            for i, step in enumerate(next_steps, 1):
                step_para = self.doc.add_paragraph()
                step_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = step_para.add_run(f"{i}. {step}")
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                run.font.color.rgb = MSLogoColors.DARK_GRAY

            self.doc.add_paragraph()  # spacer

        # Contact Information
        contact_heading = self.doc.add_paragraph()
        contact_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = contact_heading.add_run(f"{Icons.EMAIL}  Contact Information")
        run.font.name = 'Segoe UI'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = MSLogoColors.BLUE

        self.doc.add_paragraph()

        # Primary contact card
        self._add_contact_card(contact_name, contact_email, contact_role,
                               contact_team, github_usernames)

        # Additional contacts
        if additional_contacts:
            for contact in additional_contacts:
                self._add_contact_card(
                    contact.get("name", ""),
                    contact.get("email", ""),
                    contact.get("role", ""),
                    contact.get("team", contact_team),
                    contact.get("github", [])
                )

        # Spacer
        for _ in range(2):
            self.doc.add_paragraph()

        # Bottom 4-color bar
        self.add_color_bar()

        # Confidentiality
        conf = self.doc.add_paragraph()
        conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = conf.add_run("Microsoft Confidential")
        run.font.name = 'Segoe UI'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = MSLogoColors.DARK_RED

    def _add_contact_card(self, name, email, role, team, github=None):
        """Add a formatted contact card with optional GitHub usernames."""
        # Name
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(name)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = MSLogoColors.DARK_GRAY
        name_para.paragraph_format.space_after = Pt(2)

        # Role
        role_para = self.doc.add_paragraph()
        role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = role_para.add_run(role)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.color.rgb = MSLogoColors.MEDIUM_GRAY
        role_para.paragraph_format.space_before = Pt(0)
        role_para.paragraph_format.space_after = Pt(2)

        # Team
        team_para = self.doc.add_paragraph()
        team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = team_para.add_run(team)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.color.rgb = MSLogoColors.MEDIUM_GRAY
        team_para.paragraph_format.space_before = Pt(0)
        team_para.paragraph_format.space_after = Pt(2)

        # Email
        email_para = self.doc.add_paragraph()
        email_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = email_para.add_run(f"{Icons.EMAIL}  {email}")
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.color.rgb = MSLogoColors.BLUE
        email_para.paragraph_format.space_before = Pt(0)
        email_para.paragraph_format.space_after = Pt(2)

        # GitHub usernames
        if github:
            gh_text = "  ".join(github) if isinstance(github, list) else github
            gh_para = self.doc.add_paragraph()
            gh_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # GitHub icon (octocat unicode or text)
            icon_run = gh_para.add_run("GitHub: ")
            icon_run.font.name = 'Segoe UI'
            icon_run.font.size = Pt(11)
            icon_run.font.bold = True
            icon_run.font.color.rgb = MSLogoColors.DARK_GRAY

            user_run = gh_para.add_run(gh_text)
            user_run.font.name = 'Segoe UI'
            user_run.font.size = Pt(11)
            user_run.font.color.rgb = MSLogoColors.DARK_GRAY

            gh_para.paragraph_format.space_before = Pt(0)
            gh_para.paragraph_format.space_after = Pt(12)

    # =========================================================================
    # FULL DOCUMENT CREATION (Convenience)
    # =========================================================================
    def create_full_document(self):
        """
        Create complete document structure:
        1. Cover page
        2. Table of Contents
        3. Header/Footer setup

        Call this first, then add content, then call add_closing_page() last.
        """
        self.create_cover_page()
        self.add_table_of_contents()
        self.setup_header_footer()

    # =========================================================================
    # UTILITY METHODS
    # =========================================================================
    def add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()

    def add_horizontal_line(self):
        """Add a horizontal line separator."""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(12)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _set_cell_background(self, cell, color_hex: str):
        """Set table cell background color."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    def _remove_table_borders(self, table):
        """Remove all borders from a table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            borders.append(border)
        tblPr.append(borders)

    def _add_page_number(self, paragraph):
        """Add Page X of Y to paragraph."""
        page_run = paragraph.add_run("Page ")
        page_run.font.name = 'Segoe UI'
        page_run.font.size = Pt(8)
        page_run.font.color.rgb = MSLogoColors.GRAY

        # Current page
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText1 = OxmlElement('w:instrText')
        instrText1.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run = paragraph.add_run()
        run._r.append(fldChar1)
        run._r.append(instrText1)
        run._r.append(fldChar2)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(8)
        run.font.color.rgb = MSLogoColors.GRAY

        of_run = paragraph.add_run(" of ")
        of_run.font.name = 'Segoe UI'
        of_run.font.size = Pt(8)
        of_run.font.color.rgb = MSLogoColors.GRAY

        # Total pages
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.text = "NUMPAGES"
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        run2 = paragraph.add_run()
        run2._r.append(fldChar3)
        run2._r.append(instrText2)
        run2._r.append(fldChar4)
        run2.font.name = 'Segoe UI'
        run2.font.size = Pt(8)
        run2.font.color.rgb = MSLogoColors.GRAY

    # =========================================================================
    # SAVE
    # =========================================================================
    def save(self, filename: str = None):
        """
        Save the document.

        Args:
            filename: Output path. Auto-generates if not provided.
        """
        if not filename:
            filename = self.generate_filename()
        elif not filename.endswith('.docx'):
            filename += '.docx'
        self.doc.save(filename)
        print(f"✅ Document saved: {filename}")
        return filename

    def generate_filename(self) -> str:
        """Generate filename: {snake_case_title}_v{version}.docx"""
        name = self.title.lower()
        name = re.sub(r'[\s\-]+', '_', name)
        name = re.sub(r'[^a-z0-9_]', '', name)
        name = re.sub(r'_+', '_', name).strip('_')
        return f"{name}_v{self.version}.docx"


# =============================================================================
# CONVENIENCE FUNCTION
# =============================================================================
def create_colorful_document(
    title: str,
    subtitle: str = "",
    version: str = "1.0",
    client: str = "",
    author: str = "paulasilva@microsoft.com",
    subject: str = "",
    keywords: str = "",
    category: str = "Technical Documentation",
    logo_path: str = None,
    filename: str = None
) -> MSColorfulDocCreator:
    """
    Quick function: create a colorful document with full structure
    (cover with logo + TOC + headers/footers).

    Returns MSColorfulDocCreator instance ready for content.
    Call add_closing_page() when done, then save().
    """
    doc = MSColorfulDocCreator(
        title=title,
        subtitle=subtitle,
        version=version,
        client=client,
        author=author,
        subject=subject,
        keywords=keywords,
        category=category,
        logo_path=logo_path,
    )
    doc.create_full_document()

    if filename:
        doc.save(filename)

    return doc


# =============================================================================
# MAIN - Example
# =============================================================================
if __name__ == "__main__":
    doc = create_colorful_document(
        title="Agentic DevOps Workshop",
        subtitle="Accelerating AI-Native Software Delivery",
        version="1.0",
        client="Enterprise Client",
        subject="Agentic DevOps, GitHub Copilot, Azure",
        keywords="DevOps, AI, GitHub, Azure, Workshop",
        category="Workshop Guide",
    )

    doc.add_executive_summary(
        "This workshop guides enterprise teams through adopting AI-native "
        "software delivery practices using GitHub Copilot and Azure."
    )

    doc.add_color_bar()
    doc.add_section_heading("Prerequisites", color="green")
    doc.add_bullets([
        "GitHub Enterprise Cloud account",
        "Azure subscription with contributor access",
        "VS Code with GitHub Copilot extension",
    ])

    doc.add_color_bar()
    doc.add_section_heading("Agenda", color="yellow")
    doc.add_table(
        headers=["Time", "Topic", "Type"],
        rows=[
            ["9:00 AM", "Introduction & AI Maturity", "Presentation"],
            ["10:30 AM", "GitHub Copilot Deep Dive", "Hands-on Lab"],
            ["12:00 PM", "Lunch Break", "—"],
            ["1:00 PM", "Agentic DevOps Pipeline", "Hands-on Lab"],
            ["3:00 PM", "Next Steps & Wrap-up", "Discussion"],
        ],
        color="yellow"
    )

    doc.add_color_bar()
    doc.add_section_heading("Key Takeaways", color="red")
    doc.add_key_takeaway(
        "AI-native DevOps reduces cycle time by 40% on average when "
        "properly implemented with GitHub Copilot Agent Mode."
    )
    doc.add_best_practice(
        "Start with L1 (AI-Assisted) before jumping to L3 (AI-Native). "
        "Each maturity level builds on the previous one."
    )

    doc.add_closing_page(
        next_steps=[
            "Schedule follow-up POC session",
            "Define success criteria for AI adoption",
            "Set up GitHub Copilot trial licenses",
        ]
    )

    doc.save()